# noScribe - AI-powered Audio Transcription
# Copyright (C) 2023 Kai Dröge

# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.

# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.

# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.


import tkinter as tk
import customtkinter as ctk
from tkHyperlinkManager import HyperlinkManager
import webbrowser
from functools import partial
from PIL import Image
import os
import platform
import yaml
import locale
import appdirs
from subprocess import run, Popen, PIPE, STDOUT
if platform.system() == 'Windows':
    from subprocess import STARTUPINFO, STARTF_USESHOWWINDOW
from docx import Document
import docx
import re
# from pyannote.audio import Pipeline (> imported on demand below)
from typing import Any, Mapping, Optional, Text
import sys
from itertools import islice
from threading import Thread
from queue import Queue, Empty
from tempfile import TemporaryDirectory
import datetime
from contextlib import redirect_stderr, redirect_stdout
from pathlib import Path
from io import StringIO
from elevate import elevate
if platform.system() == 'Windows':
    import cpufeature
if platform.system() == "Darwin": # = MAC
    import shlex

if platform.system() == "Darwin": # = MAC
    bundle_dir = os.path.abspath(os.path.dirname(__file__))
    # if platform.machine() == "arm64": # Intel should also support MPS
    if platform.mac_ver()[0] >= '12.3': # MPS needs macOS 12.3+
        os.environ["PYTORCH_ENABLE_MPS_FALLBACK"] = str(1)

app_version = '0.3'
ctk.set_appearance_mode('dark')
ctk.set_default_color_theme('blue')

# config
config_dir = appdirs.user_config_dir('noScribe')
if not os.path.exists(config_dir):
    os.makedirs(config_dir)
try:
    with open(f'{config_dir}/config.yml', 'r') as file:
        config = yaml.safe_load(file)
except: # seems we run it for the first time and there is no config file
    config = {}

def save_config():
    with open(f'{config_dir}\\config.yml', 'w') as file:
        yaml.safe_dump(config, file)

# locale: setting the language of the UI
# see https://pypi.org/project/python-i18n/
import i18n
from i18n import t
i18n.set('filename_format', '{locale}.{format}')
if platform.system() == 'Windows':
    i18n.load_path.append('./trans')
elif platform.system() == "Darwin": # = MAC
    i18n.load_path.append(os.path.join(bundle_dir, 'trans'))
else:
    raise Exception('Platform not supported yet.')
try:
    app_locale = config['locale']
except:
    app_locale = 'auto'

if app_locale == 'auto': # read system locale settings
    try:
        app_locale = locale.getdefaultlocale()[0][0:2]
    except:
        app_locale = 'en'
i18n.set('fallback', 'en')
i18n.set('locale', app_locale)

# Check CPU capabilities and select the right version of whisper
if platform.system() == 'Windows':
    if cpufeature.CPUFeature["AVX2"] == True and cpufeature.CPUFeature["OS_AVX"] == True:
        whisper_path = "./whisper_avx2"
    else:
        whisper_path = "./whisper_sse2"
elif platform.system() == "Darwin": # = MAC
    if platform.machine() == "arm64":
        whisper_path = os.path.join(bundle_dir, "whisper_mac_arm64")
    elif platform.machine() == "x86_64":
        whisper_path = os.path.join(bundle_dir, "whisper_mac_x86_64")
    else:
        raise Exception('Could not detect Apple architecture.')
else:
    raise Exception('Platform not supported yet.')

# timestamp regex
timestamp_re = re.compile('\[\d\d:\d\d:\d\d.\d\d\d --> \d\d:\d\d:\d\d.\d\d\d\]')

# Helper functions

def millisec(timeStr): # convert 'hh:mm:ss' string to milliseconds
    try:
        spl = timeStr.split(':')
        s = (int)((int(spl[0]) * 60 * 60 + int(spl[1]) * 60 + float(spl[2]) )* 1000)
        return s
    except:
        raise Exception(t('err_invalid_time_string', time = timeStr))

def iter_except(function, exception):
        # Works like builtin 2-argument `iter()`, but stops on `exception`.
        try:
            while True:
                yield function()
        except exception:
            return
        
def docx_add_bookmark(first_run, last_run, bookmark_name, bookmark_id):
    # adds a bookmark including the two runs and everything inbetween 
    # bookmark_id must be unique
    start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
    start.set(docx.oxml.ns.qn('w:id'), str(bookmark_id))
    start.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    first_run._r.append(start)

    end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
    end.set(docx.oxml.ns.qn('w:id'), str(bookmark_id))
    end.set(docx.oxml.ns.qn('w:name'), bookmark_name)
    last_run._r.append(end)



class TimeEntry(ctk.CTkEntry): # special Entry box to enter time in the format hh:mm:ss
                               # based on https://stackoverflow.com/questions/63622880/how-to-make-python-automatically-put-colon-in-the-format-of-time-hhmmss
    def __init__(self, master, **kwargs):
        ctk.CTkEntry.__init__(self, master, **kwargs)
        vcmd = self.register(self.validate)

        self.bind('<Key>', self.format)
        self.configure(validate="all", validatecommand=(vcmd, '%P'))

        self.valid = re.compile('^\d{0,2}(:\d{0,2}(:\d{0,2})?)?$', re.I)

    def validate(self, text):
        if text == '':
            return True
        elif ''.join(text.split(':')).isnumeric():
            return not self.valid.match(text) is None
        else:
            return False

    def format(self, event):
        if event.keysym not in ['BackSpace', 'Shift_L', 'Shift_R', 'Control_L', 'Control_R']:
            i = self.index('insert')
            if i in [2, 5]:
                if event.char != ':':
                    if self.get()[i:i+1] != ':':
                        self.insert(i, ':')

class App(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.protocol("WM_DELETE_WINDOW", self.on_closing)

        self.audio_file = ''
        self.transcript_file = ''
        self.log_file = None
        self.cancel = False # if set to True, transcription will be canceled

        # configure window
        self.title('noScribe - ' + t('app_header'))
        self.geometry(f"{1100}x{650}")
        self.iconbitmap('noScribeLogo.ico')

        # header
        self.frame_header = ctk.CTkFrame(self, height=100)
        self.frame_header.pack(padx=0, pady=0, anchor='nw', fill='x')

        self.frame_header_logo = ctk.CTkFrame(self.frame_header, fg_color='transparent')
        self.frame_header_logo.pack(anchor='w', side='left')

        # logo
        self.logo_label = ctk.CTkLabel(self.frame_header_logo, text="noScribe", font=ctk.CTkFont(size=42, weight="bold"))
        self.logo_label.pack(padx=20, pady=[40, 0], anchor='w')

        # sub header
        self.header_label = ctk.CTkLabel(self.frame_header_logo, text=t('app_header'), font=ctk.CTkFont(size=16, weight="bold"))
        self.header_label.pack(padx=20, pady=[0, 20], anchor='w')
     
        # graphic
        if platform.system() == 'Windows':
            self.header_graphic = ctk.CTkImage(dark_image=Image.open('graphic_sw.png'), size=(926,119))
        elif platform.system() == "Darwin": # = MAC
            self.header_graphic = ctk.CTkImage(dark_image=Image.open(os.path.join(bundle_dir, 'graphic_sw.png')), size=(926,119))
        else:
            raise Exception('Platform not supported yet.')
        self.header_graphic_label = ctk.CTkLabel(self.frame_header, image=self.header_graphic, text='')
        self.header_graphic_label.pack(anchor='ne', side='right', padx=[30,30])


        # main window
        self.frame_main = ctk.CTkFrame(self)
        self.frame_main.pack(padx=0, pady=0, anchor='nw', expand=True, fill='both')

        # create sidebar frame for options
        self.sidebar_frame = ctk.CTkFrame(self.frame_main, width=270, corner_radius=0, fg_color='transparent')
        self.sidebar_frame.pack(padx=0, pady=0, fill='y', expand=False, side='left')

        # input audio file
        self.label_audio_file = ctk.CTkLabel(self.sidebar_frame, text=t('label_audio_file'))
        self.label_audio_file.pack(padx=20, pady=[20,0], anchor='w')

        self.frame_audio_file = ctk.CTkFrame(self.sidebar_frame, width=250, height=33, corner_radius=8, border_width=2)
        self.frame_audio_file.pack(padx=20, pady=[0,10], anchor='w')
        
        self.label_audio_file_name = ctk.CTkLabel(self.frame_audio_file, width=200, corner_radius=8, anchor='w', text=t('label_audio_file_name'))
        self.label_audio_file_name.place(x=3, y=3)

        self.button_audio_file = ctk.CTkButton(self.frame_audio_file, width=45, height=29, text='📂', command=self.button_audio_file_event)
        self.button_audio_file.place(x=203, y=2)

        # input transcript file name
        self.label_transcript_file = ctk.CTkLabel(self.sidebar_frame, text=t('label_transcript_file'))
        self.label_transcript_file.pack(padx=20, pady=[10,0], anchor='w')

        self.frame_transcript_file = ctk.CTkFrame(self.sidebar_frame, width=250, height=33, corner_radius=8, border_width=2)
        self.frame_transcript_file.pack(padx=20, pady=[0,10], anchor='w')
        
        self.label_transcript_file_name = ctk.CTkLabel(self.frame_transcript_file, width=200, corner_radius=8, anchor='w', text=t('label_transcript_file_name'))
        self.label_transcript_file_name.place(x=3, y=3)

        self.button_transcript_file = ctk.CTkButton(self.frame_transcript_file, width=45, height=29, text='📂', command=self.button_transcript_file_event)
        self.button_transcript_file.place(x=203, y=2)

        # Options grid
        self.frame_options = ctk.CTkFrame(self.sidebar_frame, width=250, fg_color='transparent')
        self.frame_options.pack(padx=20, pady=10, anchor='w', fill='x')

        self.frame_options.grid_columnconfigure(0, weight=1)
        self.frame_options.grid_columnconfigure(1, weight=0)

        # Start/stop
        self.label_start = ctk.CTkLabel(self.frame_options, text=t('label_start'))
        self.label_start.grid(column=0, row=0, sticky='w', pady=[0,5])

        self.entry_start = TimeEntry(self.frame_options, width=100)
        self.entry_start.grid(column='1', row='0', sticky='e', pady=[0,5])
        self.entry_start.insert(0, '00:00:00')

        self.label_stop = ctk.CTkLabel(self.frame_options, text=t('label_stop'))
        self.label_stop.grid(column=0, row=1, sticky='w', pady=[5,10])

        self.entry_stop = TimeEntry(self.frame_options, width=100)
        self.entry_stop.grid(column='1', row='1', sticky='e', pady=[5,10])
    
        # language
        self.label_language = ctk.CTkLabel(self.frame_options, text=t('label_language'))
        self.label_language.grid(column=0, row=2, sticky='w', pady=10)

        self.langs = ('auto', 'en (english)', 'zh (chinese)', 'de (german)', 'es (spanish)', 'ru (russian)', 'ko (korean)', 'fr (french)', 'ja (japanese)', 'pt (portuguese)', 'tr (turkish)', 'pl (polish)', 'ca (catalan)', 'nl (dutch)', 'ar (arabic)', 'sv (swedish)', 'it (italian)', 'id (indonesian)', 'hi (hindi)', 'fi (finnish)', 'vi (vietnamese)', 'iw (hebrew)', 'uk (ukrainian)', 'el (greek)', 'ms (malay)', 'cs (czech)', 'ro (romanian)', 'da (danish)', 'hu (hungarian)', 'ta (tamil)', 'no (norwegian)', 'th (thai)', 'ur (urdu)', 'hr (croatian)', 'bg (bulgarian)', 'lt (lithuanian)', 'la (latin)', 'mi (maori)', 'ml (malayalam)', 'cy (welsh)', 'sk (slovak)', 'te (telugu)', 'fa (persian)', 'lv (latvian)', 'bn (bengali)', 'sr (serbian)', 'az (azerbaijani)', 'sl (slovenian)', 'kn (kannada)', 'et (estonian)', 'mk (macedonian)', 'br (breton)', 'eu (basque)', 'is (icelandic)', 'hy (armenian)', 'ne (nepali)', 'mn (mongolian)', 'bs (bosnian)', 'kk (kazakh)', 'sq (albanian)', 'sw (swahili)', 'gl (galician)', 'mr (marathi)', 'pa (punjabi)', 'si (sinhala)', 'km (khmer)', 'sn (shona)', 'yo (yoruba)', 'so (somali)', 'af (afrikaans)', 'oc (occitan)', 'ka (georgian)', 'be (belarusian)', 'tg (tajik)', 'sd (sindhi)', 'gu (gujarati)', 'am (amharic)', 'yi (yiddish)', 'lo (lao)', 'uz (uzbek)', 'fo (faroese)', 'ht (haitian   creole)', 'ps (pashto)', 'tk (turkmen)', 'nn (nynorsk)', 'mt (maltese)', 'sa (sanskrit)', 'lb (luxembourgish)', 'my (myanmar)', 'bo (tibetan)', 'tl (tagalog)', 'mg (malagasy)', 'as (assamese)', 'tt (tatar)', 'haw (hawaiian)', 'ln (lingala)', 'ha (hausa)', 'ba (bashkir)', 'jw (javanese)', 'su (sundanese)')

        self.option_menu_language = ctk.CTkOptionMenu(self.frame_options, width=100, values=self.langs)
        self.option_menu_language.grid(column=1, row=2, sticky='e', pady=10)
        try:
            self.option_menu_language.set(config['last_language'])
        except:
            pass

        # Speaker Detection (Diarization)
        self.label_speaker = ctk.CTkLabel(self.frame_options, text=t('label_speaker'))
        self.label_speaker.grid(column=0, row=3, sticky='w', pady=10)

        self.option_menu_speaker = ctk.CTkOptionMenu(self.frame_options, width=100, values=['auto', 'none'])
        self.option_menu_speaker.grid(column=1, row=3, sticky='e', pady=10)
        try:
            self.option_menu_speaker.set(config['last_speaker'])
        except:
            pass

        # Quality (Model Selection)
        self.label_quality = ctk.CTkLabel(self.frame_options, text=t('label_quality'))
        self.label_quality.grid(column=0, row=4, sticky='w', pady=10)
        
        self.option_menu_quality = ctk.CTkOptionMenu(self.frame_options, width=100, values=['precise', 'fast'])
        self.option_menu_quality.grid(column=1, row=4, sticky='e', pady=10)
        try:
            self.option_menu_quality.set(config['last_quality'])
        except:
            pass

        # Start Button
        self.start_button = ctk.CTkButton(self.sidebar_frame, height=42, text=t('start_button'), command=self.button_start_event)
        self.start_button.pack(padx=20, pady=[0,10], expand=True, fill='x', anchor='sw')

        # Stop Button
        self.stop_button = ctk.CTkButton(self.sidebar_frame, height=42, fg_color='darkred', hover_color='darkred', text=t('stop_button'), command=self.button_stop_event)
    
        # create log textbox
        self.log_frame = ctk.CTkFrame(self.frame_main, corner_radius=0, fg_color='transparent')
        self.log_frame.pack(padx=0, pady=0, fill='both', expand=True, side='right')

        self.log_textbox = ctk.CTkTextbox(self.log_frame, wrap='word', state="disabled", font=("",16), text_color="lightgray")
        self.log_textbox.tag_config('highlight', foreground='darkorange')
        self.log_textbox.tag_config('error', foreground='yellow')
        self.log_textbox.pack(padx=20, pady=[20,10], expand=True, fill='both')

        self.hyperlink = HyperlinkManager(self.log_textbox._textbox)
        
        # status bar bottom
        self.frame_status = ctk.CTkFrame(self, height=20, corner_radius=0)
        self.frame_status.pack(padx=0, pady=[0,0], anchor='sw', fill='x', side='bottom')

        self.progress_bar = ctk.CTkProgressBar(self.frame_status, height=5, mode='determinate')
        self.progress_bar.set(0)
        
        self.logn(t('welcome_message'), 'highlight')
        self.log(t('welcome_credits', v=app_version))
        self.logn('https://github.com/kaixxx/noScribe', link='https://github.com/kaixxx/noScribe#readme')
        self.logn(t('welcome_instructions'))       
        
    # Events and Methods
    
    def openLink(self, link):
        webbrowser.open(link)

    def log(self, txt='', tags=[], where='both', link=''): # log to main window (log can be 'screen', 'file' or 'both')
        if where != 'file':
            self.log_textbox.configure(state=ctk.NORMAL)
            if link != '':
                tags = tags + self.hyperlink.add(partial(self.openLink, link))
            self.log_textbox.insert(ctk.END, txt, tags)
            self.log_textbox.yview_moveto(1) # scroll to last line
            self.log_textbox.configure(state=ctk.DISABLED)
        if (where != 'screen') and (self.log_file != None) and (self.log_file.closed == False):
            if tags == 'error':
                txt = f'ERROR: {txt}'
            self.log_file.write(txt)
    
    def logn(self, txt='', tags=[], where='both', link=''): # log with newline
        self.log(f'{txt}\n', tags, where, link)

    def logr(self, txt='', tags=[], where='both', link=''): # replace the last line of the log
        if where != 'file':
            self.log_textbox.configure(state=ctk.NORMAL)
            self.log_textbox.delete('end-2l linestart', 'end-1l')
        self.logn(txt, tags, where, link)
        
    def reader_thread(self, q):
        try:
            with self.process.stdout as pipe:
                for line in iter(pipe.readline, b''):
                    q.put(line)
        finally:
            q.put(None)

    def button_audio_file_event(self):
        fn = tk.filedialog.askopenfilename()
        if fn != '':
            self.audio_file = fn
            self.logn(t('log_audio_file_selected') + self.audio_file)
            self.label_audio_file_name.configure(text=os.path.basename(self.audio_file))

    def button_transcript_file_event(self):
        fn = tk.filedialog.asksaveasfilename(filetypes=[('Microsoft Word','*.docm')], defaultextension='docm')
        if fn != '':
            self.transcript_file = fn
            self.logn(t('log_transcript_filename') + self.transcript_file)
            self.label_transcript_file_name.configure(text=os.path.basename(self.transcript_file))
    
    def set_progress(self, step, value):
        if step == 1:
            self.progress_bar.set(value * 0.05 / 100)
        elif step == 2:
            progr = 0.05 # (step 1)
            progr = progr + (value * 0.45 / 100)
            self.progress_bar.set(progr)
        elif step == 3:
            if self.speaker_detection == 'auto':
                progr = 0.05 + 0.45 # (step 1 + step 2)
                progr_factor = 0.5
            else:
                progr = 0.05 # (step 1)
                progr_factor = 0.95
            progr = progr + (value * progr_factor / 100)
            self.progress_bar.set(progr)
        else:
            self.progress_bar.set(0)
        self.update()


    ################################################################################################
    # main function Button Start

    def button_start_event(self):
        
        proc_start_time = datetime.datetime.now()
        self.cancel = False

        # Show the stop button
        self.start_button.pack_forget() # hide
        self.stop_button.pack(padx=20, pady=[0,10], expand=True, fill='x', anchor='sw')
        
        # Show the progress bar
        self.progress_bar.set(0)
        self.progress_bar.pack(padx=20, pady=[10,20], expand=True, fill='both')

        tmpdir = TemporaryDirectory('noScribe')
        self.tmp_audio_file = tmpdir.name + '/' + 'tmp_audio.wav'
     
        try:
            # collect all the options
            if self.audio_file == '':
                self.logn(t('err_no_audio_file'), 'error')
                tk.messagebox.showerror(title='noScribe', message=t('err_no_audio_file'))
                return
            
            if self.transcript_file == '':
                self.logn(t('err_no_transcript_file'), 'error')
                tk.messagebox.showerror(title='noScribe', message=t('err_no_transcript_file'))
                return

            self.my_transcript_file = self.transcript_file

            val = self.entry_start.get()
            if val == '':
                self.start = 0
            else:
                self.start = millisec(val)
            
            val = self.entry_stop.get()
            if val == '':
                self.stop = '0'
            else:
                self.stop = millisec(val)
            
            if self.option_menu_quality.get() == 'fast':
                try:
                    self.whisper_model = config['model_path_fast']
                except:
                    if platform.system() == 'Windows':
                        config['model_path_fast'] = './models/ggml-small.bin'
                    elif platform.system() == "Darwin": # = MAC
                        config['model_path_fast'] = os.path.join(bundle_dir, 'models', 'ggml-small.bin')
                    else:
                        raise Exception('Platform not supported yet.')
                    self.whisper_model = config['model_path_fast']
            else:
                try:
                    self.whisper_model = config['model_path_precise']
                except:
                    if platform.system() == 'Windows':
                        config['model_path_precise'] = './models/ggml-large.bin'
                    elif platform.system() == "Darwin": # = MAC
                        config['model_path_fast'] = os.path.join(bundle_dir, 'models', 'ggml-large.bin')
                    else:
                        raise Exception('Platform not supported yet.')
                    self.whisper_model = config['model_path_precise']

            self.prompt = ''
            try:
                with open('prompt.yml', 'r') as file:
                    prompts = yaml.safe_load(file)
            except:
                prompts = {}

            self.language = self.option_menu_language.get()
            if self.language != 'auto':
                self.language = self.language[0:3].strip()
                try:
                    self.prompt = prompts[self.language]
                except:
                    self.prompt = ''
            
            self.speaker_detection = self.option_menu_speaker.get()

            try:
                if config['auto_save'] == 'True': # auto save during transcription (every 20 sec)?
                    self.auto_save = True
                else:
                    self.auto_save = False
            except:
                config['auto_save'] = 'True'
                self.auto_save = True 

            # create log file
            if not os.path.exists(f'{config_dir}/log'):
                os.makedirs(f'{config_dir}/log')
            self.log_file = open(f'{config_dir}/log/{Path(self.audio_file).stem}.log', 'w', encoding="utf-8")

            # log CPU capabilities
            if platform.system() == 'Windows':
                self.logn("=== CPU FEATURES ===", where="file")
                for key, value in cpufeature.CPUFeature.items():
                    self.logn('    {:24}: {}'.format(key, value), where="file")
            
            try:

                #-------------------------------------------------------
                # 1) Convert Audio

                try:
                    self.logn()
                    self.logn(t('start_audio_conversion'), 'highlight')
                    self.update()
                
                    if int(self.stop) > 0: # transcribe only part of the audio
                        end_pos_cmd = f'-to {self.stop}ms'
                    else: # tranbscribe until the end
                        end_pos_cmd = ''

                    if platform.system() == 'Windows':
                        ffmpeg_cmd = f'ffmpeg.exe -loglevel warning -y -ss {self.start}ms {end_pos_cmd} -i \"{self.audio_file}\" -ar 16000 -ac 1 -c:a pcm_s16le {self.tmp_audio_file}'
                    elif platform.system() == "Darwin":  # = MAC
                        ffmpeg_abspath = os.path.join(bundle_dir, 'ffmpeg')
                        ffmpeg_cmd = f'{ffmpeg_abspath} -nostdin -loglevel warning -y -ss {self.start}ms {end_pos_cmd} -i \"{self.audio_file}\" -ar 16000 -ac 1 -c:a pcm_s16le {self.tmp_audio_file}'
                        ffmpeg_cmd = shlex.split(ffmpeg_cmd)
                    else:
                        raise Exception('Platform not supported yet.')
                    self.logn(ffmpeg_cmd, where='file')

                    if platform.system() == 'Windows':
                        # (supresses the terminal, see: https://stackoverflow.com/questions/1813872/running-a-process-in-pythonw-with-popen-without-a-console)
                        startupinfo = STARTUPINFO()
                        startupinfo.dwFlags |= STARTF_USESHOWWINDOW
                        with Popen(ffmpeg_cmd, stdout=PIPE, stderr=STDOUT, bufsize=1,universal_newlines=True,encoding='utf-8', startupinfo=startupinfo) as ffmpeg_proc:
                            for line in ffmpeg_proc.stdout:
                                self.logn('ffmpeg: ' + line)
                    elif platform.system() == "Darwin":  # = MAC
                        with Popen(ffmpeg_cmd, stdout=PIPE, stderr=STDOUT, bufsize=1,universal_newlines=True,encoding='utf-8') as ffmpeg_proc:
                            for line in ffmpeg_proc.stdout:
                                self.logn('ffmpeg: ' + line)
                    if ffmpeg_proc.returncode > 0:
                        raise Exception(t('err_ffmpeg'))
                    self.logn(t('audio_conversion_finished'))
                    self.set_progress(1, 50)
                except Exception as e:
                    self.logn(t('err_converting_audio'), 'error')
                    self.logn(e, 'error')
                    return

                #-------------------------------------------------------
                # 2) Speaker identification (diarization) with pyannote
                
                # Helper Functions:

                def overlap_len(ss_start, ss_end, ts_start, ts_end):
                    # ss...: speaker section start and end in milliseconds (from pyannote)
                    # ts...: transcript section start and end (from whisper.cpp)
                    if ts_end < ss_start: # no overlap, ts is before ss
                        return -1   
                    elif ts_start > ss_end: # no overlap, ts is after ss
                        return 0
                    else: # ss & ts have overlap
                        if ts_start > ss_start: # ts starts after ss
                            overlap_start = ts_start
                        else:
                            overlap_start = ss_start
                        if ts_end > ss_end: # ts ends after ss
                            overlap_end = ss_end
                        else:
                            overlap_end = ts_end
                        return overlap_end - overlap_start + 1

                def find_speaker(diarization, transcript_start, transcript_end):
                    # Looks for the segment in diarization that has the most overlap with section_start-end. 
                    # Returns the speaker name if found, an empty string otherwise
                    spkr = ''
                    overlap = 0
                    
                    for segment, _, label in diarization.itertracks(yield_label=True):
                        t = overlap_len(int(segment.start * 1000), int((segment.start + segment.duration) * 1000), transcript_start, transcript_end)
                        if t == -1: # we are already after transcript_end
                            break
                        elif t > overlap:
                            overlap = t
                            spkr = f'S{label[8:]}' # shorten the label: "SPEAKER_01" > "S01"
                    return spkr

                class SimpleProgressHook:
                    #Hook to show progress of each internal step
                    def __init__(self, parent, transient: bool = False):
                        super().__init__()
                        self.parent = parent
                        self.transient = transient

                    def __enter__(self):
                        self.progress = 0
                        return self

                    def __exit__(self, *args):
                        pass
                        # self.parent.logn() # print the final new line

                    def __call__(
                        self,
                        step_name: Text,
                        step_artifact: Any,
                        file: Optional[Mapping] = None,
                        total: Optional[int] = None,
                        completed: Optional[int] = None,
                    ):        
                        # check for unser cancelation
                        if self.parent.cancel == True:
                            raise Exception(t('err_user_cancelation')) 
                        
                        if completed is None:
                            completed = total = 1

                        if not hasattr(self, 'step_name') or step_name != self.step_name:
                            self.step_name = step_name
                        
                        progress_percent = int(completed/total*100)
                        self.parent.logr(f'{step_name}: {progress_percent}%')
                        
                        if self.step_name == 'segmentation':
                            self.parent.set_progress(2, progress_percent * 0.3)
                        elif self.step_name == 'embeddings':
                            self.parent.set_progress(2, 30 + (progress_percent * 0.7))
                        
                        self.parent.update()

                # Start Diarization:

                if self.speaker_detection == 'auto':
                    try: 
                        with redirect_stderr(StringIO()) as f:
                            self.logn()
                            self.logn(t('start_identifiying_speakers'), 'highlight')
                            self.logn(t('loading_pyannote'))
                            self.update()
                            from pyannote.audio import Pipeline # import only on demand because this library is huge
                            self.set_progress(1, 100)

                            if platform.system() == 'Windows':
                                pipeline = Pipeline.from_pretrained('./models/pyannote_config.yaml')
                            elif platform.system() == "Darwin": # = MAC
                                with open(os.path.join(bundle_dir, 'models', 'pyannote_config.yaml'), 'r') as yaml_file:
                                    pyannote_config = yaml.safe_load(yaml_file)

                                pyannote_config['pipeline']['params']['embedding'] = os.path.join(bundle_dir, *pyannote_config['pipeline']['params']['embedding'].split("/")[1:])
                                pyannote_config['pipeline']['params']['segmentation'] = os.path.join(bundle_dir, *pyannote_config['pipeline']['params']['segmentation'].split("/")[1:])

                                with open(os.path.join(bundle_dir, 'models', 'pyannote_config_macOS.yaml'), 'w') as yaml_file:
                                    yaml.safe_dump(pyannote_config, yaml_file)

                                pipeline = Pipeline.from_pretrained(os.path.join(bundle_dir, 'models', 'pyannote_config_macOS.yaml'))
                                # if platform.machine() == "arm64": # Intel should also support MPS
                                if platform.mac_ver()[0] >= '12.3': # MPS needs macOS 12.3+
                                    pipeline.to("mps")
                            else:
                                raise Exception('Platform not supported yet.')
                            self.logn()
                            with SimpleProgressHook(parent=self) as hook:
                                diarization = pipeline(self.tmp_audio_file, hook=hook) # apply the pipeline to the audio file

                            # write segments to log file 
                            for segment, _, label in diarization.itertracks(yield_label=True):
                                line = (
                                    f'{int(segment.start * 1000)} {int((segment.start + segment.duration) * 1000)} {label}\n'
                                )
                                self.log(line, where='file')
                                
                            self.logn()
                            
                            # read stderr and log it:
                            err = f.readline()
                            while err != '':
                                self.logn(err, 'error')
                                err = f.readline()
                    except Exception as e:
                        self.logn(t('err_identifying_speakers'), 'error')
                        self.logn(e, 'error')
                        return

                #-------------------------------------------------------
                # 3) Transcribe with whisper.cpp

                self.logn()
                self.logn(t('start_transcription'), 'highlight')
                self.logn(t('loading_whisper'))
                self.logn()
                self.update()
                
                # prompt?
                if self.prompt != '':
                    self.prompt_cmd = f'--prompt "{self.prompt}"'
                else:
                    self.prompt_cmd = ''
                
                # whisper options:
                try:
                    # max segement length. Shorter segments can improve speaker identification.
                    self.whisper_options = f"--max-len {config['whisper_options_max-len']}" 
                except:
                    config['whisper_options_max-len'] = '30'
                    self.whisper_options = "--max-len 30"
                
                # "whisper_extra_commands" can be defined in config.yml and will be attached to the end of the command line. 
                # Use this to experiment with advanced options.
                # see https://github.com/ggerganov/whisper.cpp/tree/master/examples/main for a list of options
                # Be careful: If your options change the output of main.exe in the terminal, noScribe might not be able to interpret this and fail badly...

                try:
                    self.whisper_extra_commands = config['whisper_extra_commands']
                    if self.whisper_extra_commands == None:
                        self.whisper_extra_commands = ''
                except:
                    config['whisper_extra_commands'] = ''
                    self.whisper_extra_commands = ''
                
                command = f'{whisper_path}/main --model {self.whisper_model} --language {self.language} {self.prompt_cmd} {self.whisper_options} --print-colors --print-progress --file "{self.tmp_audio_file}" {self.whisper_extra_commands}'
                if platform.system() == "Darwin":  # = MAC
                    command = shlex.split(command)
                self.logn(command, where='file')

                # prepare transcript docm
                if platform.system() == 'Windows':
                    d = Document('transcriptTempl.docm')
                elif platform.system() == "Darwin": # = MAC
                    d = Document(os.path.join(bundle_dir,'transcriptTempl.docm'))
                else:
                    raise Exception('Platform not supported yet.')
                d.core_properties.author = f'noScribe vers. {app_version}'
                d.core_properties.comments = self.audio_file
                
                # header
                p = d.paragraphs[0]
                p.text = Path(self.audio_file).stem # use the name of the audio file (without extension) as the title
                p.style = 'noScribe_header'
                
                p = d.add_paragraph(t('doc_header', version=app_version), style='noScribe_subheader')
                p = d.add_paragraph(t('doc_header_audio', file=self.audio_file), style='noScribe_subheader')
                
                p = d.add_paragraph()
                speaker = ''
                bookmark_id = 0
                self.last_auto_save = datetime.datetime.now()

                def save_doc():
                    try:
                        d.save(self.my_transcript_file)
                        self.last_auto_save = datetime.datetime.now()
                    except:
                        # saving failed, maybe the file is already open in Word and cannot be overwritten
                        # try saving to a different filename
                        transcript_path = Path(self.my_transcript_file)
                        self.my_transcript_file = f'{transcript_path.parent}/{transcript_path.stem}_1.docm'
                        if os.path.exists(self.my_transcript_file):
                            # the alternative filename also exists already, don't want to overwrite, giving up
                            raise Exception(t('rescue_saving_failed'))
                        else:
                            d.save(self.my_transcript_file)
                            self.logn()
                            self.logn(t('rescue_saving', file=self.my_transcript_file), 'error')
                            self.last_auto_save = datetime.datetime.now()
            
                try:
                    if platform.system() == 'Windows':
                        startupinfo = STARTUPINFO()
                        startupinfo.dwFlags |= STARTF_USESHOWWINDOW
                        self.process = Popen(command, stdout=PIPE, stderr=STDOUT, startupinfo=startupinfo)
                    elif platform.system() == "Darwin":  # = MAC
                        self.process = Popen(command, stdout=PIPE, stderr=STDOUT)
                    # Run whisper.cpp main.exe without blocking the GUI:
                    # Source: https://stackoverflow.com/questions/12057794/python-using-popen-poll-on-background-process 
                    # launch thread to read the subprocess output
                    #   (put the subprocess output into the queue in a background thread,
                    #    get output from the queue in the GUI thread.
                    #    Output chain: process.readline -> queue -> GUI)
                    q = Queue(maxsize=1024)  # limit output buffering (may stall subprocess)
                    th = Thread(target=self.reader_thread, args=[q])
                    th.daemon = True # close pipe if GUI process exits
                    th.start()

                    while self.process.poll() == None: # process is running
                        self.update()
                        # check for unser cancelation
                        if self.cancel == True:
                            if self.auto_save == True:
                                save_doc()
                                self.logn()
                                self.logn(t('transcription_saved', file=self.my_transcript_file))
                                raise Exception(t('err_user_cancelation')) 
                            else:    
                                raise Exception(t('err_user_cancelation')) 
                        # process lines from the queue
                        for line in iter_except(q.get_nowait, Empty):
                            if line is None:
                                break
                            else:
                                line = str(line.decode("utf-8", errors='ignore')) # convert to regular string
                                
                                # check if we have a transcript line from stdout or a line from stdterr
                                if timestamp_re.match(line) != None: 
                                    # found a timestamp, must be a transcript
                                    
                                    line = line.replace('\n', '') # remove line breaks     
                                    line = line.replace('\r', '') # remove carriage return     

                                    # get time of the segment in milliseconds
                                    #[00:00:00.000 --> 00:00:05.760]    
                                    start = line[1:13]
                                    end = line[18:30]
                                    start = millisec(start)
                                    end = millisec(end)
                                    
                                    line = line[33:] # discard timestamp
                                    line = line.lstrip() # discard leading spaces
                                    
                                    # write text to the doc
                                    # diarization (speaker detection)?
                                    if self.speaker_detection == 'auto':
                                        spkr = find_speaker(diarization, start, end)
                                        if (speaker != spkr) & (spkr != ''):
                                            speaker = spkr
                                            self.logn()
                                            p = d.add_paragraph()
                                            line = f'{speaker}: {line}'

                                    first_run = p.add_run() # empty run for start_bookmark
                                    # check for confidence level markers (colors)
                                    if line.find('\u001B[38;5;') > -1: 
                                        line_segments = line.split('\u001B[38;5;')
                                        cl_markers = {'196m': 1, '202m': 2, '208m': 3, '214m': 4, '220m': 5, '226m': 6, '190m': 7, '154m': 8, '118m': 9, '82m': 10}
                                        for s in line_segments:
                                            if s == '':
                                                continue # skip empty segments
                                            # extract confidence level marker, get the level from cl_markers: 
                                            cl_marker_end = s.find('m')
                                            if cl_marker_end in [2,3]: # only possible positions
                                                cl_marker = s[0:cl_marker_end + 1]
                                                if cl_marker in cl_markers:
                                                    cl_level = cl_markers[cl_marker]
                                                else: # invalid marker
                                                    cl_level = 0
                                            else: # marker not found
                                                cl_level = 0
                                            # add segments to doc
                                            s = s[cl_marker_end + 1:] # delete cl_marker
                                            s = s.replace('\u001B[0m', '') # delete the closing cl mark 
                                            r = p.add_run()
                                            r.text = s
                                            # Mark confidence level with a character based style,'noScribe_cl[1-10]'
                                            # This way, we can color-mark specific levels later in Word.
                                            if cl_level > 0:
                                                r.style = d.styles[f'noScribe_cl{cl_level}']
                                            self.log(s)
                                    else: # no marker in line
                                            r = p.add_run()
                                            r.text = line
                                            self.log(line)
                                    
                                    # Create bookmark with audio timestamps start to end.
                                    # This way, we can jump to the according audio position and play it later in Word.
                                    bookmark_id = bookmark_id + 1
                                    last_run = p.add_run()
                                    # if we skipped a part at the beginning of the audio we have to add this here again, otherwise the timestaps will not match the original audio:
                                    orig_audio_start = self.start + start
                                    orig_audio_end = self.start + end
                                    docx_add_bookmark(first_run, last_run, f'ts_{orig_audio_start}_{orig_audio_end}', bookmark_id)
                                    
                                    # auto save
                                    if self.auto_save == True:
                                        if (datetime.datetime.now() - self.last_auto_save).total_seconds() > 20:
                                            save_doc()    

                                    self.update()

                                else: # must be line from stderr
                                    self.logn(line, where='file')
                                    if line[0:35] == 'whisper_full_with_state: progress =': # progress
                                        progr = int(''.join(filter(str.isdigit, line))) # extract number, see https://stackoverflow.com/questions/4289331/how-to-extract-numbers-from-a-string-in-python
                                        self.set_progress(3, progr)

                    save_doc()
                    self.logn()
                    self.logn()
                    self.logn(t('transcription_finished'), 'highlight')
                    if self.transcript_file != self.my_transcript_file: # used alternative filename because saving under the initial name failed
                        self.logn(t('rescue_saving', file=self.my_transcript_file), 'error')
                    else:
                        self.log(t('transcription_saved'))
                        self.logn(self.my_transcript_file, link=f'file://{self.my_transcript_file}')
                    # log duration of the whole process in minutes
                    proc_time = datetime.datetime.now() - proc_start_time
                    self.logn(t('trancription_time', duration=int(proc_time.total_seconds() / 60))) 

                    if self.process.poll() > 0:
                        raise Exception(t('err_whisper_main', e=self.process.poll()))
                
                except Exception as e:
                    self.logn()
                    self.logn(t('err_transcription'), 'error')
                    self.logn(e, 'error')
                    return
                
                finally:
                    self.process.kill() # exit subprocess (zombie!)
            finally:
                self.log_file.close()
                self.log_file = None

        except Exception as e:
            self.logn(t('err_options'), 'error')
            self.logn(e, 'error')
            return
        
        finally:
            # hide the stop button
            self.stop_button.pack_forget() # hide
            self.start_button.pack(padx=20, pady=[0,10], expand=True, fill='x', anchor='sw')

            # hide progress bar
            self.progress_bar.pack_forget()

    # End main function Button Start        
    ################################################################################################

    def button_stop_event(self):
        if tk.messagebox.askyesno(title='noScribe', message=t('transcription_canceled')) == True:
            self.logn()
            self.logn(t('start_canceling'))
            self.update()
            self.cancel = True

    def on_closing(self):
        # (see: https://stackoverflow.com/questions/111155/how-do-i-handle-the-window-close-event-in-tkinter)
        #if messagebox.askokcancel("Quit", "Do you want to quit?"):
        try:
            # remember some settings for the next run
            config['last_language'] = self.option_menu_language.get()
            config['last_speaker'] = self.option_menu_speaker.get()
            config['last_quality'] = self.option_menu_quality.get()
            save_config()
        finally:
            self.destroy()

if __name__ == "__main__":

    app = App()
    
    app.mainloop()